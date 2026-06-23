import os
import re
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
import fitz

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_shading(cell, color):
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{color}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def set_table_all_borders(table):
    tblPr = table._tbl.tblPr
    borders_xml = (
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(parse_xml(borders_xml))

def clear_table_borders(table):
    tblPr = table._tbl.tblPr
    borders_xml = (
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="none"/>'
        f'<w:bottom w:val="none"/>'
        f'<w:left w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'<w:insideH w:val="none"/>'
        f'<w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(parse_xml(borders_xml))

def set_col_widths(table, widths):
    for i, col in enumerate(table.columns):
        col.width = widths[i]
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = widths[i]

def configure_header(doc):
    for section in doc.sections:
        # Define margens (em polegadas: 1 in = 2.54 cm)
        section.top_margin = Inches(1.5)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.header_distance = Inches(0.4)
        
        header = section.header
        # Limpa parágrafos padrões do cabeçalho
        for p in header.paragraphs:
            p.text = ""
            
        # Cria uma tabela no cabeçalho para alinhar os logos nas extremidades
        # Largura total disponível: 6.5 polegadas (8.5 largura total - 1.0 esq - 1.0 dir)
        header_table = header.add_table(1, 2, Inches(6.5))
        clear_table_borders(header_table)
        header_table.columns[0].width = Inches(3.25)
        header_table.columns[1].width = Inches(3.25)
        
        # Logo da esquerda (UniCesumar)
        cell_left = header_table.cell(0, 0)
        set_cell_margins(cell_left, top=0, bottom=0, left=0, right=0)
        p_left = cell_left.paragraphs[0]
        p_left.paragraph_format.space_after = Pt(0)
        r_left = p_left.add_run()
        r_left.add_picture('Documentação final - Mybuddy/header_logo_left.png', width=Inches(1.6))
        
        # Logo da direita (Dog)
        cell_right = header_table.cell(0, 1)
        set_cell_margins(cell_right, top=0, bottom=0, left=0, right=0)
        p_right = cell_right.paragraphs[0]
        p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_right.paragraph_format.space_after = Pt(0)
        r_right = p_right.add_run()
        r_right.add_picture('Documentação final - Mybuddy/header_logo_right.jpg', width=Inches(0.65))

def parse_requisitos_pdf():
    pdf_path = "Documentação final - Mybuddy/documentações diversas/MyBuddy - Requisitos F e NF.pdf"
    if not os.path.exists(pdf_path):
        raise FileNotFoundError(f"Arquivo não encontrado: {pdf_path}")
        
    doc = fitz.open(pdf_path)
    full_text = ""
    for page in doc:
        full_text += page.get_text() + "\n"
        
    # Extração de RFs
    rf_pattern = re.compile(r'(RF-\d{3})\s+(UC\d{2})?\s+(.*?)\s+(Alta|Média|Baixa)', re.DOTALL)
    rf_matches = rf_pattern.findall(full_text)
    rf_list = []
    for rf_id, uc, desc, prio in rf_matches:
        desc_clean = " ".join([w.strip() for w in desc.split() if w.strip()])
        rf_list.append((rf_id, uc or "-", desc_clean, prio))
        
    # Extração de RNs
    rn_pattern = re.compile(r'(RN-\d{3})\s+(.*?)(?=RN-\d{3}|\b(?:5\.\s+Requisitos|UniCesumar)\b|\Z)', re.DOTALL)
    rn_matches = rn_pattern.findall(full_text)
    rn_list = []
    for rn_id, desc in rn_matches:
        desc_clean = " ".join([w.strip() for w in desc.split() if w.strip()])
        rn_list.append((rn_id, desc_clean))
        
    # Extração de RNFs com normalização de quebra de linhas nas palavras chave
    norm_text = re.sub(r'RNF-0\s*\n\s*(\d{2})', r'RNF-0\1', full_text)
    norm_text = re.sub(r'Desempenh\s*\n\s*o', 'Desempenho', norm_text)
    norm_text = re.sub(r'Seguran\s*\n\s*ça', 'Segurança', norm_text)
    norm_text = re.sub(r'Disponibilid\s*\n\s*ade', 'Disponibilidade', norm_text)
    norm_text = re.sub(r'Manutenibili\s*\n\s*dade', 'Manutenibilidade', norm_text)
    norm_text = re.sub(r'Escalabilida\s*\n\s*de', 'Escalabilidade', norm_text)
    
    rnf_pattern = re.compile(r'(RNF-0\d{2})\s+(\w+)\s+(.*?)(?=(?:RNF-0\d{2})|\b(?:6\.\s+Critérios|UniCesumar|7\.\s+Conclusão)\b|\Z)', re.DOTALL)
    rnf_matches = rnf_pattern.findall(norm_text)
    rnf_list = []
    for rnf_id, cat, desc in rnf_matches:
        desc_clean = " ".join([w.strip() for w in desc.split() if w.strip()])
        rnf_list.append((rnf_id, cat, desc_clean))
        
    return rf_list, rn_list, rnf_list

def main():
    print("Iniciando extração do PDF de Requisitos...")
    rf_list, rn_list, rnf_list = parse_requisitos_pdf()
    print(f"Extraídos: {len(rf_list)} RFs, {len(rn_list)} RNs, {len(rnf_list)} RNFs.")
    
    doc = docx.Document()
    
    # Configura fontes globais
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Configuração de Cabeçalho / Geometria de Página
    configure_header(doc)
    
    COLOR_PRIMARY = RGBColor(0, 0, 0)
    COLOR_SECONDARY = RGBColor(0, 0, 0)
    
    # --- Estilos Auxiliares ---
    def add_heading_1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = COLOR_PRIMARY
        return p

    def add_heading_2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = COLOR_SECONDARY
        return p

    def add_body_paragraph(text):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(10)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(11)
        return p

    def add_bullet_point(text, bold_prefix=""):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(4)
        if bold_prefix:
            run_prefix = p.add_run(bold_prefix)
            run_prefix.font.name = "Arial"
            run_prefix.font.bold = True
            run_prefix.font.size = Pt(11)
        run_text = p.add_run(text)
        run_text.font.name = "Arial"
        run_text.font.size = Pt(11)
        return p

    # --- PÁGINA 1: CAPA ---
    # Adiciona espaçamento vertical para centralizar o título
    p_spacer = doc.add_paragraph()
    p_spacer.paragraph_format.space_before = Pt(36)
    
    capa_header = doc.add_paragraph()
    capa_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_h = capa_header.add_run(
        "Análise e Desenvolvimento de Sistemas\n"
        "Gestão de Projetos Tecnológicos\n"
        "4º. semestre - Noturno"
    )
    run_h.font.name = "Arial"
    run_h.font.size = Pt(12)
    run_h.font.bold = True
    
    p_spacer2 = doc.add_paragraph()
    p_spacer2.paragraph_format.space_before = Pt(48)
    
    capa_integrantes = doc.add_paragraph()
    capa_integrantes.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_i = capa_integrantes.add_run(
        "Eder Henrique Pontes  R.A 24534211-2\n\n"
        "Julia Cardoso R.A 24503170-2\n\n"
        "Davi Cassoli Lira R.A 24042075-2\n\n"
        "Daniel Godinho R.A 24383624-2"
    )
    run_i.font.name = "Arial"
    run_i.font.size = Pt(11)
    run_i.font.bold = True

    p_spacer3 = doc.add_paragraph()
    p_spacer3.paragraph_format.space_before = Pt(24)

    # Big Dog Illustration in center
    capa_image = doc.add_paragraph()
    capa_image.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_img = capa_image.add_run()
    r_img.add_picture('Documentação final - Mybuddy/header_logo_right.jpg', width=Inches(3.2))

    p_spacer4 = doc.add_paragraph()
    p_spacer4.paragraph_format.space_before = Pt(18)

    capa_footer = doc.add_paragraph()
    capa_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_f = capa_footer.add_run("Maringá\n2025")
    run_f.font.name = "Arial"
    run_f.font.size = Pt(11)
    run_f.font.bold = True

    doc.add_page_break()

    # --- PÁGINA 2: SUMÁRIO ---
    sumario_title = doc.add_paragraph()
    sumario_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_sum = sumario_title.add_run("Sumário")
    run_sum.font.name = "Arial"
    run_sum.font.size = Pt(14)
    run_sum.font.bold = True
    
    doc.add_paragraph("\n")
    
    sum_items = [
        "1. VISÃO GERAL E FUNDAMENTOS DO PRODUTO .......................................................................... 3",
        "   1.1 Contexto e Justificativa ......................................................................................................... 3",
        "   1.2 Metodologia de Descoberta (Double Diamond) .......................................................................... 3",
        "   1.3 Análise de Mercado (SWOT) ................................................................................................... 4",
        "   1.4 Business Model Canvas .......................................................................................................... 4",
        "   1.5 Personas e Mapa de Empatia ................................................................................................... 5",
        "   1.6 Glossário dos Termos do Negócio .............................................................................................. 6",
        "2. ENGENHARIA DE REQUISITOS ..................................................................................................... 7",
        "   2.1 Atores do Sistema ................................................................................................................... 7",
        "   2.2 Requisitos Funcionais (RF) ......................................................................................................... 7",
        "   2.3 Requisitos Não Funcionais (RNF) ................................................................................................... 8",
        "   2.4 Regras de Negócio e LGPD ......................................................................................................... 9",
        "3. SYSTEM DESIGN E ARQUITETURA DE SOFTWARE .................................................................... 10",
        "   3.1 Arquitetura Base (Spring Boot + Angular + Flutter) .......................................................................... 10",
        "   3.2 Estratégia Dual Database (PostgreSQL + MongoDB) ...................................................................... 10",
        "   3.3 Cenário de Alta Escalabilidade ......................................................................................................... 10",
        "   3.4 Comunicação em Tempo Real (WebSocket) ..................................................................................... 11",
        "4. MODELAGEM UML E ORIENTAÇÃO A OBJETOS .............................................................................. 12",
        "   4.1 Diagrama de Casos de Uso ........................................................................................................... 12",
        "   4.2 Diagrama de Sequência .................................................................................................................. 12",
        "   4.3 Diagrama de Classes e Entidade-Relacionamento .............................................................................. 12",
        "5. EXPERIÊNCIA DO USUÁRIO (UX/UI) E FRONTENDS ......................................................................... 13",
        "   5.1 Design System e Protótipos (Figma) ............................................................................................ 13",
        "   5.2 Aplicação Web (Angular) ............................................................................................................. 13",
        "   5.3 Aplicativo Multiplataforma (Flutter) ................................................................................................ 13",
        "6. INTEGRAÇÕES DE TERCEIROS, SEGURANÇA E DEVOPS ................................................................ 14",
        "   6.1 Gestão de Identidade (Keycloak) ................................................................................................. 14",
        "   6.2 Motor Financeiro (Mercado Pago Split) ........................................................................................... 14",
        "   6.3 Infraestrutura, Docker e Deploy ....................................................................................................... 14",
        "7. GARANTIA DE QUALIDADE (QA) E TESTES ................................................................................... 15",
        "   7.1 Critérios de Aceitação do MVP ......................................................................................................... 15",
        "   7.2 Testes Automatizados Implementados ........................................................................................... 15",
        "   7.3 Documentação da API REST (Swagger) .......................................................................................... 16",
        "8. GESTÃO, PROCESSOS ÁGEIS E GOVERNANÇA ............................................................................... 17",
        "   8.1 Metodologia Kanban e Jira ............................................................................................................. 17",
        "   8.2 Gestão de Configuração e Commits (GitHub) .................................................................................. 17",
        "   8.3 Apontamento de Horas e Roadmap Executado ............................................................................... 17",
        "REFERÊNCIAS BIBLIOGRÁFICAS ......................................................................................................... 18",
        "ANEXOS E CANAIS DO PROJETO ......................................................................................................... 19"
    ]
    
    for item in sum_items:
        p_item = doc.add_paragraph()
        p_item.paragraph_format.line_spacing = 1.15
        p_item.paragraph_format.space_after = Pt(3)
        run_item = p_item.add_run(item)
        run_item.font.name = "Arial"
        run_item.font.size = Pt(10)
        
    doc.add_page_break()

    # --- CAPÍTULO 1: VISÃO GERAL E FUNDAMENTOS ---
    add_heading_1("1. VISÃO GERAL E FUNDAMENTOS DO PRODUTO")
    
    add_heading_2("1.1 Contexto e Justificativa")
    add_body_paragraph(
        "Este documento detalha o projeto MyBuddy, uma plataforma web Restful concebida para servir como "
        "um ecossistema digital centralizado para a causa animal em nível local. O projeto visa otimizar "
        "o processo de adoção de animais de estimação, conectando ONGs e protetores a potenciais adotantes "
        "de forma eficiente e segura. Adicionalmente, a plataforma funcionará como um guia de serviços pet, "
        "oferecendo aos tutores um acesso facilitado a informações sobre pet shops e clínicas veterinárias "
        "em sua região, fortalecendo a comunidade e promovendo a posse responsável."
    )
    add_body_paragraph(
        "A necessidade do projeto MyBuddy surge de um problema central: a fragmentação do processo de adoção "
        "e da busca por serviços para pets. Atualmente, ONGs e protetores lutam por visibilidade em redes "
        "sociais desordenadas, onde suas publicações se perdem em meio a outros conteúdos, dificultando o alcance "
        "de um público qualificado. Do outro lado, potenciais adotantes enfrentam uma busca frustrante e desorganizada, "
        "gerando desconfiança e ineficiência. A falta de uma fonte centralizada e confiável é o principal obstáculo. "
        "O MyBuddy justifica-se pela oportunidade de resolver essa dor, oferecendo uma solução 'dois em um' que "
        "unifica a jornada de 'adotar' e 'cuidar', algo raro no mercado, e atendendo a uma clara preferência do público "
        "por plataformas digitais para gerenciar a vida de seus pets."
    )

    add_heading_2("1.2 Metodologia de Descoberta (Double Diamond)")
    add_body_paragraph(
        "O processo de design do MyBuddy seguiu as quatro fases do Duplo Diamante:\n"
        "1. Descobrir: Foco na ampla coleta de informações através de pesquisas com adotantes, entrevistas com protetores "
        "e análise de concorrentes. O problema identificado foi a conexão fragmentada e ineficiente entre quem deseja adotar "
        "e quem tem um animal para doação.\n"
        "2. Definir: Síntese dos dados para focar no problema central: a falta de uma fonte centralizada e confiável. O foco "
        "do projeto foi definido como atuar como um hub local para conectar ONGs, adotantes e serviços.\n"
        "3. Desenvolver: Geração de ideias e soluções através de brainstorming de funcionalidades e criação de protótipos "
        "de baixa fidelidade das telas principais.\n"
        "4. Entregar: Teste dos protótipos com usuários reais, implementação das funcionalidades priorizadas com base no "
        "feedback e lançamento da primeira versão do sistema (MVP) para um novo ciclo de melhorias."
    )

    add_heading_2("1.3 Análise de Mercado (SWOT)")
    add_body_paragraph(
        "O mercado pet no Brasil é um dos mais relevantes do mundo, com projeção de faturamento superior a R$ 70 bilhões "
        "em 2025. O principal vetor de crescimento é a 'humanização' dos pets, vistos como membros da família, o que "
        "aumenta a disposição dos tutores para investir em serviços de qualidade e optar pela adoção responsável. A "
        "jornada do tutor é cada vez mais digital, criando um ambiente fértil para soluções como o MyBuddy. A concorrência "
        "se divide em plataformas de adoção de foco nacional, canais alternativos desorganizados e agregadores de serviços "
        "não especializados. A principal oportunidade para o MyBuddy é se posicionar como o hub da comunidade pet local, "
        "unificando a jornada de adoção e cuidado, uma brecha deixada pelos concorrentes. A análise SWOT detalhada resume "
        "esses pontos a seguir:"
    )

    # Tabela SWOT
    t_swot = doc.add_table(rows=1, cols=4)
    set_table_all_borders(t_swot)
    set_col_widths(t_swot, [Inches(1.625)] * 4)
    
    # Cabeçalho SWOT
    swot_headers = ["Forças", "Fraquezas", "Oportunidades", "Ameaças"]
    for i, h_text in enumerate(swot_headers):
        cell = t_swot.rows[0].cells[i]
        set_cell_shading(cell, "F2F2F2")
        set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(h_text)
        run.font.name = "Arial"
        run.font.bold = True
        run.font.size = Pt(10)
        
    swot_row = t_swot.add_row()
    swot_contents = [
        "• Foco Hiperlocal, criando senso de comunidade.\n\n• Solução Integrada (adotar e cuidar), algo raro no mercado.\n\n• Gratuito para ONGs, o que facilita a adesão inicial.",
        "• Marca desconhecida, sem confiança inicial.\n\n• Recursos de marketing e desenvolvimento limitados.\n\n• Necessidade de atrair tanto animais quanto adotantes simultaneamente.",
        "• Crescimento contínuo do mercado pet.\n\n• Frustração dos usuários com a desorganização das redes sociais.\n\n• Possibilidade de parcerias com comércios e prefeituras locais.",
        "• Concorrência estabelecida que pode copiar funcionalidades.\n\n• Baixa adesão inicial pode tornar a plataforma irrelevante.\n\n• Risco de informações desatualizadas minarem a confiança do usuário."
    ]
    for i, content in enumerate(swot_contents):
        cell = swot_row.cells[i]
        set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(content)
        run.font.name = "Arial"
        run.font.size = Pt(9.5)

    doc.add_paragraph("\n")

    add_heading_2("1.4 Business Model Canvas")
    add_body_paragraph(
        "A fim de estruturar o modelo de negócios e garantir viabilidade comercial, desenvolveu-se o "
        "Business Model Canvas do MyBuddy, detalhando como a plataforma gera e entrega valor à comunidade pet:"
    )

    # Tabela Business Model Canvas (9 linhas, 4 colunas)
    t_canvas = doc.add_table(rows=1, cols=4)
    set_table_all_borders(t_canvas)
    set_col_widths(t_canvas, [Inches(1.2), Inches(1.76), Inches(1.76), Inches(1.76)])
    
    # Cabeçalho Canvas
    canvas_headers = ["Bloco", "Detalhamento A", "Detalhamento B", "Detalhamento C"]
    for i, h_text in enumerate(canvas_headers):
        cell = t_canvas.rows[0].cells[i]
        set_cell_shading(cell, "F2F2F2")
        set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(h_text)
        run.font.name = "Arial"
        run.font.bold = True
        run.font.size = Pt(9.5)

    canvas_rows = [
        ("Proposta de Valor", 
         "Para Adotantes e Tutores: Uma plataforma centralizada e confiável que simplifica a busca por um novo pet e por serviços locais essenciais, como pet shops e veterinários.",
         "Para ONGs e Protetores: Um canal de divulgação gratuito, poderoso e organizado para aumentar significativamente a visibilidade dos animais e suas chances de adoção.",
         "Para a Comunidade Pet Local: Um hub digital que conecta todos os pontos do ecossistema animal, fortalecendo o comércio local e promovendo a posse responsável."),
         
        ("Segmentos de Clientes",
         "Usuários Finais (Adotantes/Tutores): Pessoas e famílias que buscam adotar um animal de estimação ou que já possuem um e precisam de serviços locais.",
         "Fornecedores de Conteúdo (ONGs e Protetores): Organizações e indivíduos que cuidam de animais e os disponibilizam para adoção, sendo a fonte principal de conteúdo.",
         "Parceiros de Serviço (Pet Shops e Veterinários): Estabelecimentos e profissionais que desejam divulgar seus serviços para um público altamente segmentado."),
         
        ("Canais",
         "Plataforma Web Principal: O sistema web responsivo (Desktop e Mobile) é o principal canal de interação e transações.",
         "Redes Sociais: Utilizadas ativamente para engajamento comunitário, marketing e divulgação de animais e parceiros.",
         "Parcerias e Mídia Local: Colaborações com comércios, prefeitura e imprensa local para ampliar o alcance e a credibilidade."),
         
        ("Relacionamento com Clientes",
         "Plataforma Self-Service: A interface será intuitiva e de fácil navegação, permitindo que a maioria dos usuários se sirva de forma autônoma.",
         "Construção de Comunidade: Foco em criar um senso de comunidade através do compartilhamento de histórias de sucesso e do foco no ecossistema local.",
         "Suporte Direto: Ofereceremos canais de suporte para parceiros-chave, como ONGs e patrocinadores."),
         
        ("Fontes de Receita",
         "Foco Inicial em Crescimento: A plataforma será mantida gratuitamente no início para construir uma base sólida de usuários e validar o modelo.",
         "Monetização Futura: Potenciais fontes de receita incluem patrocínios de empresas locais, planos de destaque para parceiros de serviço e doações da comunidade.",
         "Publicidade Contextual: Inserção de anúncios de produtos ou serviços relevantes para o público pet, exibidos de forma não intrusiva."),
         
        ("Atividades-Chave",
         "Desenvolvimento e Manutenção: Construir, manter e evoluir o software é a atividade técnica central da equipe.",
         "Marketing e Comunidade: Atrair novos usuários (adotantes e protetores) e manter a comunidade ativa.",
         "Gestão de Parcerias: Cadastrar e dar suporte às ONGs, protetores e parceiros de serviço para garantir conteúdo de qualidade."),
         
        ("Recursos-Chave",
         "A Plataforma MyBuddy: O software e os dados hospedados são os principais ativos intelectuais do projeto.",
         "Equipe de Desenvolvimento: A equipe técnica responsável por construir e manter a plataforma.",
         "Rede de Parceiros: A base de ONGs, protetores e estabelecimentos cadastrados é um recurso vital para a relevância."),
         
        ("Parcerias-Chave",
         "ONGs e Protetores: São os parceiros essenciais que fornecem o 'conteúdo' principal (os animais para adoção).",
         "Pet Shops e Veterinários: Parceiros comerciais que enriquecem a plataforma com o guia de serviços e marketplace.",
         "Instituições e Patrocinadores: Empresas e órgãos públicos que podem fornecer recursos financeiros ou apoio institucional."),
         
        ("Estrutura de Custos",
         "Recursos Humanos: O tempo e esforço dedicado da equipe de desenvolvimento e marketing.",
         "Infraestrutura Tecnológica: Custos com servidores de nuvem, banco de dados, domínio e ferramentas de desenvolvimento.",
         "Marketing e Divulgação: Investimentos para atrair usuários e construir a reputação da marca local.")
    ]
    
    for row_data in canvas_rows:
        row = t_canvas.add_row()
        for idx, text in enumerate(row_data):
            cell = row.cells[idx]
            set_cell_margins(cell, top=80, bottom=80, left=80, right=80)
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(text)
            run.font.name = "Arial"
            run.font.size = Pt(9)
            if idx == 0:
                run.font.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_cell_shading(cell, "F9F9F9")
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
    doc.add_page_break()

    # --- 1.5 PERSONAS E MAPA DE EMPATIA ---
    add_heading_2("1.5 Personas e Mapa de Empatia")
    add_body_paragraph(
        "Para garantir que a experiência do usuário (UX) guiasse as decisões de design de software, a equipe "
        "definiu duas Personas representativas de nosso ecossistema e estruturou um Mapa de Empatia detalhado:"
    )
    
    add_body_paragraph(
        "Persona 1 - Ana (A Adotante Planejadora):\n"
        "28 anos, Designer Gráfica. Mora em um apartamento, trabalha em modelo híbrido. É digitalmente proficiente, "
        "pesquisa muito antes de tomar decisões e busca apoiar uma causa de bem-estar animal. Quer adotar um gato de "
        "porte pequeno que se adapte bem à sua rotina de trabalho. Sente-se frustrada ao navegar por dezenas de grupos "
        "desorganizados de Facebook, onde fotos são ruins, dados são incompletos e há receio de golpes. O MyBuddy oferece "
        "a ela filtros de busca avançados, perfis detalhados e contato com protetores verificados."
    )
    
    add_body_paragraph(
        "Persona 2 - Carlos (O Protetor Apaixonado):\n"
        "45 anos, Comerciante e protetor independente. Resgata, cuida e prepara cães de rua para adoção em sua própria "
        "residência. Possui tempo limitado para divulgações e poucos recursos financeiros. Sofre com a dispersão de "
        "tempo respondendo a comentários repetitivos em redes sociais que se perdem e não atingem o público qualificado. "
        "O MyBuddy fornece a ele uma ferramenta centralizada que economiza tempo, gerencia formulários e coloca os pets "
        "em uma vitrine qualificada."
    )

    # Tabela Mapa de Empatia
    t_empatia = doc.add_table(rows=1, cols=2)
    set_table_all_borders(t_empatia)
    set_col_widths(t_empatia, [Inches(1.8), Inches(4.7)])
    
    # Cabeçalho Mapa Empatia
    cell_q = t_empatia.rows[0].cells[0]
    set_cell_shading(cell_q, "F2F2F2")
    set_cell_margins(cell_q)
    r_q = cell_q.paragraphs[0].add_run("Categoria (Persona Ana)")
    r_q.font.bold = True
    r_q.font.size = Pt(10)
    
    cell_ans = t_empatia.rows[0].cells[1]
    set_cell_shading(cell_ans, "F2F2F2")
    set_cell_margins(cell_ans)
    r_ans = cell_ans.paragraphs[0].add_run("Expressões da Persona")
    r_ans.font.bold = True
    r_ans.font.size = Pt(10)

    empatia_rows = [
        ("O que ela pensa e sente?", "Pensa: 'Eu realmente quero dar um lar para um animal, mas o processo parece tão estressante. Será que vou encontrar o pet certo para meu estilo de vida? Tenho medo de cair em algum golpe.' Sente: Ansiedade, frustração com a busca, mas também esperança e empolgação com a ideia de ter um novo companheiro."),
        ("O que ela fala e faz?", "Fala: 'Alguém conhece uma ONG de confiança?', 'Estou procurando um gatinho filhote, alguém sabe onde posso encontrar?'. Faz: Passa horas navegando em redes sociais, manda mensagens para vários contatos, visita sites de ONGs."),
        ("O que ela vê?", "Grupos de Facebook com fotos de baixa qualidade. Posts com poucas informações. Concorrentes com interfaces complexas. Histórias de sucesso de adoção de amigos e conhecidos."),
        ("O que ela ouve?", "'Adotar é a melhor coisa que você faz!'. 'Cuidado, tem muita gente que não é séria na internet'. 'Minha amiga achou o cachorro dela num grupo, mas demorou meses e teve dor de cabeça'."),
        ("Dores", "Perda de tempo em pesquisas dispersas. Insegurança e desconfiança quanto à veracidade das ONGs. Informações fragmentadas e de baixa qualidade no perfil dos animais. Medo de não encontrar um pet compatível."),
        ("Ganhos", "Encontrar o companheiro de quatro patas ideal de forma rápida. Sentir que fez a diferença de forma positiva. Ter um processo de busca simples, seguro e prazeroso. Desenvolver confiança imediata no protetor.")
    ]
    
    for cat, text in empatia_rows:
        row = t_empatia.add_row()
        
        c_cat = row.cells[0]
        set_cell_margins(c_cat, top=80, bottom=80, left=80, right=80)
        set_cell_shading(c_cat, "F9F9F9")
        p_c = c_cat.paragraphs[0]
        p_c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_c = p_c.add_run(cat)
        r_c.font.name = "Arial"
        r_c.font.bold = True
        r_c.font.size = Pt(9.5)
        
        c_txt = row.cells[1]
        set_cell_margins(c_txt, top=80, bottom=80, left=80, right=80)
        p_t = c_txt.paragraphs[0]
        p_t.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_t.paragraph_format.line_spacing = 1.15
        p_t.paragraph_format.space_after = Pt(0)
        r_t = p_t.add_run(text)
        r_t.font.name = "Arial"
        r_t.font.size = Pt(9.5)

    doc.add_paragraph("\n")

    # --- 1.6 GLOSSÁRIO ---
    add_heading_2("1.6 Glossário dos Termos do Negócio")
    add_body_paragraph(
        "Para nivelamento de domínio e cumprimento das diretrizes de engenharia de software do TCC, "
        "foram definidos os seguintes conceitos essenciais da aplicação:"
    )

    # Tabela do Glossário
    t_glos = doc.add_table(rows=1, cols=2)
    set_table_all_borders(t_glos)
    set_col_widths(t_glos, [Inches(1.8), Inches(4.7)])
    
    t_glos.rows[0].cells[0].text = "Termo / Conceito"
    t_glos.rows[0].cells[1].text = "Definição no Contexto do Projeto"
    set_cell_shading(t_glos.rows[0].cells[0], "F2F2F2")
    set_cell_shading(t_glos.rows[0].cells[1], "F2F2F2")
    set_cell_margins(t_glos.rows[0].cells[0])
    set_cell_margins(t_glos.rows[0].cells[1])
    t_glos.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
    t_glos.rows[0].cells[1].paragraphs[0].runs[0].font.bold = True
    
    glossary_terms = [
        ("Split Payment", "Divisão automática de pagamentos no checkout para repasse proporcional de comissão da plataforma e lojista."),
        ("Adotante / Tutor", "Usuário final verificado apto a preencher formulários de adoção ou gerenciar serviços de pets."),
        ("Protetor Independente", "Pessoa física que resgata e cuida de animais temporariamente sem possuir CNPJ de ONG."),
        ("Token JWT", "JSON Web Token usado para autenticação stateless, trafegado entre frontend e backend."),
        ("Dual Database", "Estratégia de banco de dados híbrida separando dados relacionais ACID (PostgreSQL) e não-relacionais (MongoDB)."),
        ("RESTful Stateless", "Padrão de API onde cada requisição contém todos os dados de autorização sem manter sessão no servidor."),
        ("WebSocket", "Protocolo de rede full-duplex de baixa latência para chat e interações síncronas."),
        ("Keycloak", "Provedor de identidade de código aberto responsável pelo controle de acessos RBAC e segurança baseada em Roles.")
    ]
    
    for term, definition in glossary_terms:
        row = t_glos.add_row()
        row.cells[0].text = term
        row.cells[1].text = definition
        set_cell_margins(row.cells[0])
        set_cell_margins(row.cells[1])
        
    for row in t_glos.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(9.5)

    doc.add_page_break()

    # --- CAPÍTULO 2: ENGENHARIA DE REQUISITOS ---
    add_heading_1("2. ENGENHARIA DE REQUISITOS")
    
    add_heading_2("2.1 Atores do Sistema")
    add_body_paragraph(
        "A especificação do escopo do MyBuddy é mapeada com base nas interações de seus atores chave. "
        "A tabela a seguir apresenta os perfis de acesso definidos na plataforma:"
    )

    t_atores = doc.add_table(rows=1, cols=2)
    set_table_all_borders(t_atores)
    set_col_widths(t_atores, [Inches(1.5), Inches(5.0)])
    
    t_atores.rows[0].cells[0].text = "Ator"
    t_atores.rows[0].cells[1].text = "Descrição e Responsabilidades"
    set_cell_shading(t_atores.rows[0].cells[0], "F2F2F2")
    set_cell_shading(t_atores.rows[0].cells[1], "F2F2F2")
    set_cell_margins(t_atores.rows[0].cells[0])
    set_cell_margins(t_atores.rows[0].cells[1])
    t_atores.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
    t_atores.rows[0].cells[1].paragraphs[0].runs[0].font.bold = True

    atores_data = [
        ("Visitante", "Acesso público à landing page, vitrine de pets (somente leitura) e guia de serviços. Não possui autenticação."),
        ("Adotante / Tutor", "Usuário autenticado que manifesta interesse em adoções, favorita pets, realiza doações, acessa o marketplace e gerencia seu próprio perfil."),
        ("ONG / Protetor", "Organização ou protetor aprovado que cadastra e gerencia pets, avalia formulários de adoção, registra acompanhamento pós-adoção e mantém a página pública da organização."),
        ("Parceiro de Serviço", "Pet shop, clínica veterinária ou hotel pet cadastrado na plataforma. Gerencia seu perfil no guia de serviços, seus produtos no marketplace e acompanha pedidos."),
        ("Lojista", "Empresa habilitada para venda online no marketplace. Gerencia inventário de produtos, promoções, estoque e status de pedidos. Habilitado via integração Mercado Pago Split."),
        ("Administrador", "Gestor da plataforma com acesso total. Aprova/rejeita cadastros de ONGs e parceiros, suspende usuários, consulta logs de auditoria, gera relatórios e configura parâmetros globais.")
    ]
    for ator, desc in atores_data:
        row = t_atores.add_row()
        row.cells[0].text = ator
        row.cells[1].text = desc
        set_cell_margins(row.cells[0])
        set_cell_margins(row.cells[1])
        
    for row in t_atores.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(9.5)

    doc.add_paragraph("\n")

    add_heading_2("2.2 Requisitos Funcionais (RF)")
    add_body_paragraph(
        "Abaixo são especificados os 52 requisitos funcionais mapeados na plataforma MyBuddy, "
        "organizados por ID, referência de Caso de Uso (UC Ref.), descrição e prioridade técnica:"
    )

    t_rf = doc.add_table(rows=1, cols=4)
    set_table_all_borders(t_rf)
    set_col_widths(t_rf, [Inches(0.8), Inches(0.8), Inches(4.1), Inches(0.8)])
    
    t_rf.rows[0].cells[0].text = "ID"
    t_rf.rows[0].cells[1].text = "UC Ref."
    t_rf.rows[0].cells[2].text = "Descrição do Requisito Funcional"
    t_rf.rows[0].cells[3].text = "Prioridade"
    
    for cell in t_rf.rows[0].cells:
        set_cell_shading(cell, "F2F2F2")
        set_cell_margins(cell, top=60, bottom=60, left=60, right=60)
        cell.paragraphs[0].runs[0].font.bold = True

    for rf_id, uc, desc, prio in rf_list:
        row = t_rf.add_row()
        row.cells[0].text = rf_id
        row.cells[1].text = uc
        row.cells[2].text = desc
        row.cells[3].text = prio
        for cell in row.cells:
            set_cell_margins(cell, top=60, bottom=60, left=60, right=60)
            
    for row in t_rf.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(8.5)

    doc.add_page_break()

    # --- 2.3 REQUISITOS NÃO FUNCIONAIS ---
    add_heading_2("2.3 Requisitos Não Funcionais (RNF)")
    add_body_paragraph(
        "Na tabela a seguir constam os 22 requisitos não funcionais organizados por categoria, "
        "detalhando as restrições técnicas, métricas de desempenho e padrões de engenharia:"
    )

    t_rnf = doc.add_table(rows=1, cols=3)
    set_table_all_borders(t_rnf)
    set_col_widths(t_rnf, [Inches(0.8), Inches(1.5), Inches(4.2)])
    
    t_rnf.rows[0].cells[0].text = "ID"
    t_rnf.rows[0].cells[1].text = "Categoria"
    t_rnf.rows[0].cells[2].text = "Descrição Técnica do Requisito"
    
    for cell in t_rnf.rows[0].cells:
        set_cell_shading(cell, "F2F2F2")
        set_cell_margins(cell, top=60, bottom=60, left=60, right=60)
        cell.paragraphs[0].runs[0].font.bold = True

    for rnf_id, cat, desc in rnf_list:
        row = t_rnf.add_row()
        row.cells[0].text = rnf_id
        row.cells[1].text = cat
        row.cells[2].text = desc
        for cell in row.cells:
            set_cell_margins(cell, top=60, bottom=60, left=60, right=60)
            
    for row in t_rnf.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(8.5)

    doc.add_page_break()

    # --- 2.4 REGRAS DE NEGÓCIO ---
    add_heading_2("2.4 Regras de Negócio e LGPD")
    add_body_paragraph(
        "As restrições de comportamento de negócios e fluxos de dados, bem como os requisitos "
        "legais de anonimização da Lei Geral de Proteção de Dados (LGPD), estão formalizados a seguir:"
    )

    t_rn = doc.add_table(rows=1, cols=2)
    set_table_all_borders(t_rn)
    set_col_widths(t_rn, [Inches(1.0), Inches(5.5)])
    
    t_rn.rows[0].cells[0].text = "ID"
    t_rn.rows[0].cells[1].text = "Descrição da Regra de Negócio"
    
    for cell in t_rn.rows[0].cells:
        set_cell_shading(cell, "F2F2F2")
        set_cell_margins(cell, top=60, bottom=60, left=60, right=60)
        cell.paragraphs[0].runs[0].font.bold = True

    for rn_id, desc in rn_list:
        row = t_rn.add_row()
        row.cells[0].text = rn_id
        row.cells[1].text = desc
        for cell in row.cells:
            set_cell_margins(cell, top=60, bottom=60, left=60, right=60)
            
    for row in t_rn.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(9.0)

    doc.add_page_break()

    # --- CAPÍTULO 3: SYSTEM DESIGN ---
    add_heading_1("3. SYSTEM DESIGN E ARQUITETURA DE SOFTWARE")
    
    add_heading_2("3.1 Arquitetura Base")
    add_body_paragraph(
        "A arquitetura do MyBuddy segue o modelo clássico cliente-servidor desacoplado com foco em "
        "APIs RESTful stateless de baixo acoplamento. O backend foi desenvolvido utilizando Java 21 com "
        "o ecossistema Spring Boot 3.5.5, aproveitando recursos modernos do framework para gestão transacional, "
        "injeção de dependências e segurança declarativa. O frontend web é baseado em Angular, aplicando o "
        "conceito de Single Page Application (SPA) para uma interface rica e fluida. No ambiente mobile, adota-se "
        "o SDK multiplataforma Flutter em linguagem Dart, garantindo compatibilidade total e alto desempenho nativo "
        "tanto no sistema operacional Android quanto no iOS."
    )

    add_heading_2("3.2 Estratégia Dual Database")
    add_body_paragraph(
        "Uma das principais decisões arquiteturais do projeto é o uso de bancos de dados híbridos (Dual Database), "
        "onde cada banco é selecionado para tratar a tipologia de dados ideal:\n"
        "• PostgreSQL Relacional: Responsável por persistir dados que exigem propriedades ACID estritas, "
        "consistência forte e relacionamentos rígidos, como transações financeiras, doações recorrentes, checkout split "
        "e agendamentos de serviços.\n"
        "• MongoDB NoSQL: Utilizado para coleções flexíveis que exigem velocidade de leitura e mudanças frequentes "
        "na modelagem, como informações cadastrais de ONGs, perfis de usuários, fotos e características detalhadas dos pets. "
        "A consistência eventual e desvinculação em deleções são gerenciadas na camada de serviço (Java backend)."
    )

    add_heading_2("3.3 Cenário de Alta Escalabilidade")
    add_body_paragraph(
        "A fim de atender às restrições do regulamento de TCC para suporte de carga extrema (1 milhão de escritas e "
        "2 milhões de leituras diárias), o design do MyBuddy incorporou:\n"
        "1. Índices Estratégicos: Aplicação da anotação @Indexed no campo petshopId e e-mail no MongoDB para acelerar buscas.\n"
        "2. Cache Redis: Feeds de pets ativos e busca por proximidade utilizam cache distribuído para evitar sobrecarga "
        "no banco principal.\n"
        "3. HikariCP Tuning: Otimização do pool de conexões do PostgreSQL e prevenção do problema de N+1 consultas com "
        "queries otimizadas JPA."
    )

    add_heading_2("3.4 Comunicação em Tempo Real")
    add_body_paragraph(
        "Para a comunicação entre adotantes e ONGs no fluxo de entrevista para adoção, implementou-se comunicação síncrona "
        "em tempo real via protocolo WebSocket. Isso elimina a latência e o overhead de rede causados por requisições HTTP "
        "periódicas (polling), permitindo um canal direto de chat síncrono e responsivo."
    )

    doc.add_page_break()

    # --- CAPÍTULO 4: MODELAGEM UML ---
    add_heading_1("4. MODELAGEM UML E ORIENTAÇÃO A OBJETOS")
    
    add_heading_2("4.1 Diagramas de Caso de Uso e Especificações")
    add_body_paragraph(
        "Foram documentados graficamente e especificados os limites do sistema e interações dos atores (Visitante, Adotante, "
        "ONG, Admin, Parceiro) com os 50 casos de uso detalhados no plano de desenvolvimento do software. "
        "As definições completas encontram-se nos arquivos Diagrama de Caso de Usos.pdf e Casos de uso - Mybuddy.xlsx."
    )

    add_heading_2("4.2 Diagramas de Sequência")
    add_body_paragraph(
        "Ilustram o fluxo de controle, troca de mensagens e ordem cronológica dos componentes para as seguintes "
        "funcionalidades críticas do sistema:\n"
        "• Autenticação integrada ao Keycloak.\n"
        "• Submissão de Formulário de Adoção Cross-DB.\n"
        "• Fluxo de Split Payment com Mercado Pago Checkout Pro.\n"
        "Consulte os fluxogramas dinâmicos detalhados no arquivo MyBuddy_DiagramasDeSequencia.pdf."
    )

    add_heading_2("4.3 Diagramas de Classes e Entidade-Relacionamento")
    add_body_paragraph(
        "Representação visual da modelagem orientada a objetos (Entidades e Relações) e a estrutura física do "
        "banco de dados PostgreSQL. O diagrama mapeia chaves estrangeiras relacionais e a lógica de atrelamento "
        "com documentos dinâmicos do MongoDB (armazenados via IDs no modelo PostgreSQL). "
        "Consulte o arquivo físico MyBuddy_DER_Final.jpg."
    )

    doc.add_page_break()

    # --- CAPÍTULO 5: EXPERIÊNCIA DO USUÁRIO ---
    add_heading_1("5. EXPERIÊNCIA DO USUÁRIO (UX/UI) E FRONTENDS")
    
    add_heading_2("5.1 Design System e Protótipos (Mockups)")
    add_body_paragraph(
        "A experiência visual do usuário do MyBuddy foi desenhada sob os pilares de acessibilidade e simplicidade "
        "operacional. A prototipação de alta fidelidade das telas web e mobile foi estruturada no Figma, definindo "
        "uma paleta coerente com tons pastéis harmônicos, tipografia moderna (Inter), micro-interações de botões "
        "e um conjunto unificado de ícones. O link oficial com o wireframe interativo está documentado em Figma.txt."
    )

    add_heading_2("5.2 Aplicação Web (Angular)")
    add_body_paragraph(
        "A interface web foi desenvolvida com Angular, aproveitando conceitos de Single Page Application (SPA), "
        "componentização de views reutilizáveis (formulários, cards de pets), rotas protegidas (Guards baseados "
        "nas Roles do Keycloak) e interceptores HTTP para adição transparente de cabeçalhos de autorização Bearer Token."
    )

    add_heading_2("5.3 Aplicativo Multiplataforma (Flutter)")
    add_body_paragraph(
        "O aplicativo mobile do MyBuddy adota o SDK Flutter em linguagem Dart. A arquitetura segue o padrão de "
        "gerenciamento de estado Cubit/BLoC para maior reatividade e injeção de dependências por construtor, "
        "garantindo paridade completa de funcionalidades operacionais em relação à versão web."
    )

    doc.add_page_break()

    # --- CAPÍTULO 6: INTEGRAÇÕES ---
    add_heading_1("6. INTEGRAÇÕES DE TERCEIROS, SEGURANÇA E DEVOPS")
    
    add_heading_2("6.1 Gestão de Identidade (Keycloak)")
    add_body_paragraph(
        "A gestão de autenticação e autorização (IAM) do MyBuddy foi delegada ao Keycloak, um servidor Identity and "
        "Access Management de nível empresarial. O Keycloak gerencia fluxos OAuth2/OIDC, tokens JWT assinados com "
        "chaves assimétricas, renovação (Refresh Tokens) e controle de acesso baseado em perfis (RBAC)."
    )

    add_heading_2("6.2 Motor Financeiro (Mercado Pago)")
    add_body_paragraph(
        "O fluxo de pagamentos de doações recorrentes e marketplace é integrado ao Mercado Pago via APIs do SDK oficial:\n"
         "• Split Payment: O valor total é rateado no momento da aprovação do pagamento, direcionando a taxa administrativa "
         "do MyBuddy e o repasse financeiro do lojista diretamente para suas respectivas carteiras digitais.\n"
         "• Webhooks: Configuração de endpoint HTTP seguro para receber notificações de alteração de status do pagamento."
    )

    add_heading_2("6.3 Infraestrutura, Docker e Deploy")
    add_body_paragraph(
        "Buscando portabilidade, isolamento de ambiente e simplificação de deploy contínuo, a aplicação foi containerizada. "
        "O arquivo Docker Compose gerencia a orquestração do PostgreSQL, MongoDB, cache Redis, Keycloak e a aplicação Java "
        "backend. O pipeline possibilita o deploy com poucos cliques em servidores como Railway ou Render."
    )

    doc.add_page_break()

    # --- CAPÍTULO 7: GARANTIA DE QUALIDADE ---
    add_heading_1("7. GARANTIA DE QUALIDADE (QA) E TESTES")
    
    add_heading_2("7.1 Critérios de Aceitação do MVP")
    add_body_paragraph(
        "Os critérios mínimos de aceitação técnica definidos para validar o MVP do TCC estão resumidos a seguir:"
    )

    # Tabela de Critérios de Aceitação
    t_mvp = doc.add_table(rows=1, cols=2)
    set_table_all_borders(t_mvp)
    set_col_widths(t_mvp, [Inches(1.8), Inches(4.7)])
    
    t_mvp.rows[0].cells[0].text = "Módulo do Sistema"
    t_mvp.rows[0].cells[1].text = "Critério de Aceitação para o MVP"
    set_cell_shading(t_mvp.rows[0].cells[0], "F2F2F2")
    set_cell_shading(t_mvp.rows[0].cells[1], "F2F2F2")
    set_cell_margins(t_mvp.rows[0].cells[0])
    set_cell_margins(t_mvp.rows[0].cells[1])
    t_mvp.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
    t_mvp.rows[0].cells[1].paragraphs[0].runs[0].font.bold = True

    mvp_data = [
        ("Autenticação e Keycloak", "Usuário deve conseguir se cadastrar, fazer login via Keycloak e receber token JWT válido. Credenciais inválidas retornam HTTP 401. Recuperação de senha funciona via link por e-mail."),
        ("Cadastro de Pet", "ONG autenticada deve cadastrar pet com fotos obrigatórias. Pet deve aparecer na vitrine paginada com todos os dados corretos."),
        ("Busca com Filtros", "Listagem deve retornar apenas registros que correspondam aos filtros aplicados (espécie, porte, cidade, estado). Paginação deve funcionar corretamente."),
        ("Fluxo Completo de Adoção", "Adotante preenche formulário de adoção → ONG visualiza a solicitação → ONG aprova ou rejeita → Adotante visualiza status atualizado → Conclusão registrada com pet marcado como Adotado."),
        ("Controle de Acesso (Roles)", "Rotas restritas por role (ONG, Admin) retornam HTTP 403 para usuários sem permissão. Admin possui acesso total ao painel de gestão."),
        ("Gestão de Organizações", "Admin consegue aprovar/rejeitar cadastro de ONG. ONG recebe e-mail com resultado. CNPJ/e-mail duplicados são rejeitados."),
        ("Cobertura de Testes", "Suíte de testes do backend atinge ≥70% de cobertura reportada pelo JaCoCo, sem falhas de compilação ou testes quebrados."),
        ("Deploy e Acesso Público", "Aplicação acessível por URL pública, com backend respondendo nos endpoints da API e frontend carregando corretamente no browser."),
        ("Conformidade LGPD", "Solicitação de exclusão de conta anonimiza corretamente os dados pessoais. Contas com processos pendentes têm a exclusão bloqueada com mensagem informativa.")
    ]
    for mod, crit in mvp_data:
        row = t_mvp.add_row()
        row.cells[0].text = mod
        row.cells[1].text = crit
        set_cell_margins(row.cells[0])
        set_cell_margins(row.cells[1])
        
    for row in t_mvp.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(9.5)

    doc.add_paragraph("\n")

    add_heading_2("7.2 Testes Automatizados Implementados")
    add_body_paragraph(
        "Para atestar o funcionamento das regras e assegurar a qualidade do software em CI/CD, a equipe implementou:\n"
        "• Testes Unitários e Integração (Backend): Utiliza JUnit 5 e Mockito para simular chamadas de serviços e validar "
        "a integridade referencial cross-DB. Adotou-se 'spring.cache.type=none' em testes para evitar problemas de poluição do contexto.\n"
        "• Testes de Interface (Frontend Angular): Cobertura de testes unitários em componentes.\n"
        "• Testes Mobile: Testagem com Dart Tests em Flutter."
    )

    add_heading_2("7.3 Documentação da API REST")
    add_body_paragraph(
        "A especificação e documentação interativa das rotas e recursos HTTP expostos pelo Spring Boot foram automatizadas "
        "com a biblioteca Springdoc OpenAPI (Swagger UI). Os endpoints detalham tipos de retorno, códigos de status "
        "e payloads de requisição/resposta."
    )

    doc.add_page_break()

    # --- CAPÍTULO 8: GESTÃO E GOVERNANÇA ---
    add_heading_1("8. GESTÃO, PROCESSOS ÁGEIS E GOVERNANÇA")
    
    add_heading_2("8.1 Metodologia Kanban")
    add_body_paragraph(
        "O ciclo de vida do desenvolvimento do MyBuddy foi orquestrado via metodologia ágil Kanban com o auxílio do Jira. "
        "O projeto foi dividido em ciclos de entregas (sprints de duas semanas) focadas na estimativa e execução de histórias "
        "de usuário (User Stories), com rituais de planejamento (Planning) e retrospectivas."
    )

    add_heading_2("8.2 Gestão de Configuração e Commits")
    add_body_paragraph(
        "A integridade da base de código no GitHub é protegida por políticas de branch. Nenhuma alteração é mesclada diretamente "
        "em Developer ou Main sem abertura prévia de Pull Request com revisão obrigatória. Os commits seguem a convenção "
        "Conventional Commits (ex: feat:, fix:, docs:, refactor:, chore:) para geração automatizada de Changelogs."
    )

    add_heading_2("8.3 Apontamento de Horas e Roadmap Executado")
    add_body_paragraph(
        "O cronograma do projeto foi dividido em três grandes épicos de desenvolvimento acadêmico, acumulando esforço real:\n"
        "• Épico 1 (Configurações, Keycloak, Base Angular): 200 horas estimadas e executadas.\n"
        "• Épico 2 (Vitrine de Pets, Adoções, APP Flutter): 340 horas planejadas.\n"
        "• Épico 3 (Marketplace, Otimizações de DB e Produção): 260 horas planejadas.\n"
        "As horas foram consolidadas individualmente nas planilhas do projeto para fins de apontamento."
    )

    doc.add_page_break()

    # --- REFERÊNCIAS ---
    add_heading_1("REFERÊNCIAS BIBLIOGRÁFICAS")
    
    refs = [
        "DEITEL, Paul; DEITEL, Harvey. Java: Como Programar. 10ª ed. São Paulo: Pearson Prentice Hall, 2016.",
        "INSTITUTO PET BRASIL (IPB). Faturamento do Setor Pet em 2024. São Paulo, 2025. Disponível em: https://institutopetbrasil.com/faturamento-do-setor-pet-2024/. Acesso em: 13 set. 2025.",
        "PRESSMAN, Roger S.; MAXIM, Bruce R. Engenharia de Software: Uma Abordagem Profissional. 8ª ed. Porto Alegre: AMGH Editora, 2016.",
        "RIES, Eric. A Startup Enxuta: Como os Empreendedores Atuais Utilizam a Inovação Contínua para Criar Empresas de Sucesso. 1ª ed. São Paulo: Leya, 2012.",
        "SEBRAE. Oportunidades no Mercado Pet: Tendências e Comportamento do Consumidor. Serviço Brasileiro de Apoio às Micro e Pequenas Empresas, 2024.",
        "SPRING. Spring Framework Documentation. Disponível em: https://spring.io/projects/spring-framework."
    ]
    
    for ref in refs:
        p_ref = doc.add_paragraph()
        p_ref.paragraph_format.line_spacing = 1.15
        p_ref.paragraph_format.space_after = Pt(8)
        p_ref.paragraph_format.left_indent = Inches(0.5)
        run_ref = p_ref.add_run(ref)
        run_ref.font.name = "Arial"
        run_ref.font.size = Pt(10)

    doc.add_paragraph("\n")

    # --- ANEXOS ---
    add_heading_1("ANEXOS E CANAIS DO PROJETO")
    
    anexos = [
        "Repositório Oficial do GitHub: https://github.com/EderHenriq/MyBuddy",
        "Quadro de Sprints no Jira: https://projetodesoftware420.atlassian.net/jira/software/projects/BUDDY/boards/34",
        "Google Drive de Artefatos e Modelagens: https://drive.google.com/drive/folders/1OEwP3-mNprn7MHA7agHBYxmT8tZzGgO2?usp=sharing",
        "Termo de Abertura do Projeto (PDF): https://drive.google.com/file/d/1IYDVTT8pxF9wcx9yYGH7cT4ZH6lUmYxk/view?usp=sharing"
    ]
    
    for anexo in anexos:
        p_an = doc.add_paragraph()
        p_an.paragraph_format.space_after = Pt(6)
        run_an = p_an.add_run(f"• {anexo}")
        run_an.font.name = "Arial"
        run_an.font.size = Pt(10)

    output_path = "Documentação final - Mybuddy/Template_de_documentação.docx"
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print("Documento final gerado e salvo com sucesso em:", output_path)

if __name__ == "__main__":
    main()
